import io
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from app.config import settings

class ExtractionError(Exception):
    """Base exception for document extraction errors."""
    pass

class UnsupportedFileTypeError(ExtractionError):
    """Raised when file extension is not supported."""
    pass

class FileSizeExceededError(ExtractionError):
    """Raised when file size exceeds the allowed limit."""
    pass

class DocumentExtractorService:
    @staticmethod
    def validate_file(filename: str, file_size_bytes: int) -> None:
        """Validates the file extension and size constraints."""
        # Validate size
        max_bytes = settings.MAX_FILE_SIZE_MB * 1024 * 1024
        if file_size_bytes > max_bytes:
            raise FileSizeExceededError(
                f"File size exceeds the maximum limit of {settings.MAX_FILE_SIZE_MB}MB."
            )

        # Validate extension
        ext = Path(filename).suffix.lower().lstrip(".")
        if ext not in settings.ALLOWED_EXTENSIONS:
            raise UnsupportedFileTypeError(
                f"File type '.{ext}' is not supported. Allowed formats: {', '.join(settings.ALLOWED_EXTENSIONS)}"
            )

    @classmethod
    def extract_text(cls, filename: str, content: bytes) -> str:
        """
        Extracts raw text from document content bytes based on file extension.
        Supports: PDF, DOCX, and TXT.
        """
        cls.validate_file(filename, len(content))
        ext = Path(filename).suffix.lower().lstrip(".")

        try:
            if ext == "txt":
                return cls._extract_txt(content)
            elif ext == "docx":
                return cls._extract_docx(content)
            elif ext == "pdf":
                return cls._extract_pdf(content)
            else:
                raise UnsupportedFileTypeError(f"Unsupported file extension: {ext}")
        except ExtractionError:
            raise
        except Exception as e:
            raise ExtractionError(f"Failed to extract text from {filename}: {str(e)}") from e

    @staticmethod
    def _extract_txt(content: bytes) -> str:
        """Extracts text from a raw TXT byte content, attempting UTF-8 then Latin-1."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return content.decode("latin-1")
            except UnicodeDecodeError as e:
                raise ExtractionError("Failed to decode TXT file with UTF-8 or Latin-1 encoding.") from e

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        """Extracts text from DOCX bytes using python-docx."""
        try:
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract from tables as well for better coverage
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                paragraphs.append(paragraph.text)

            return "\n".join(paragraphs)
        except Exception as e:
            raise ExtractionError(f"Corrupted or invalid DOCX document: {str(e)}") from e

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        """Extracts text from PDF bytes using pypdf."""
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            if reader.is_encrypted:
                try:
                    # Try decrypting with empty password
                    reader.decrypt("")
                except Exception as e:
                    raise ExtractionError("Encrypted or password-protected PDF files are not supported.") from e
            
            text_pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_pages.append(page_text)
                    
            if not text_pages:
                raise ExtractionError("No extractable text found in PDF (scanned/image-only PDFs are not supported).")
                
            return "\n\n".join(text_pages)
        except ExtractionError:
            raise
        except Exception as e:
            raise ExtractionError(f"Invalid PDF document: {str(e)}") from e
